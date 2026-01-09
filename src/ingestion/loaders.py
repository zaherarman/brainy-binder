from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from docx import Document

def load_text_file(filepath):
    """
    Load a text (.text) or a markdown (.md) file.

    Args:
        filepath: Path to the file
    Return:
        A list containing a single Document
    """
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()

    title = filepath.stem
    lines = content.split("\n")

    if lines and lines[0].startswith("#"):
        title = lines[0].lstrip("#").strip()
    
    metadata = {
        "source_path": str(filepath),
        "document_type": "note",
        "title": title,
        "filepath": filepath.suffix
    }

    return [Document(page_content=content, metadata=metadata)]

def load_pdf_file(filepath):
    """
    Load a PDF (.pdf) file.

    Args:
        filepath: Path to the pdf
    Return:
        A list containing a Document objects (usually per page)
    """
    loader = PyPDFLoader(str(filepath))

    documents = loader.load()

    title = filepath.stem

    for doc in documents:
        doc.metadata.update(
            {
                "source_path": str(filepath),
                "document_type": "pdf",
                "title": title,
                "filetype" : ".pdf"
            }
        )

    return documents

def load_word_file(filepath):
    """
    Load a Word (.docx) file.

    Args:
        filepath: Path to the word doc
    Return:
        A list containing a single Document
    """
    doc = Document(filepath)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    contents = "\n\n".join(paragraphs)

    metadata = {
        "source_path": str(filepath),
        "document_type": "word",
        "title": filepath.stem,
        "filepath": filepath.suffix
    }

    return [Document(page_content=contents, metadata=metadata)]

def load_document(filepath):
    """
    Load a document based on its file-type.

    Args:
        filepath: Path to the document

    Returns:
        List of Document objects, or None if the filetype isn't supported
    """
    suffix = filepath.suffix.lower()

    try:
        if suffix in [".txt", ".md"]:
            return load_text_file(filepath)
        elif suffix == ".pdf":
            return load_pdf_file(filepath)
        elif suffix == ".docx":
            return load_word_file(filepath)
        else:
            return None
        
    except Exception as e:
        print(f"Error loading {filepath}: {e}")
        return None

def discover_documents(data_dir):
    """
    Discover all supported documents in a directory tree.

    Args:
        data_dir: Root directory to scan

    Returns:
        List of file paths
    """
    supported_types = [".txt", ".md", ".pdf", ".docx"]
    documents = []

    for filepath in data_dir.rglob("*"):
        if filepath.is_file() and filepath.suffix.lower() in supported_types:
            documents.append(filepath)
    
    return sorted(documents)